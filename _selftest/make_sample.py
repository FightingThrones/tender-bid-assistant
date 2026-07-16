#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""自测样例生成器：造一份假标书 docx + 假公告 + 假客户画像，用于验证 pipeline。"""
import os

BASE = os.path.dirname(os.path.abspath(__file__))
SAMPLE = os.path.join(BASE, "sample.docx")
TENDER = os.path.join(BASE, "tender.txt")
COMPANY = os.path.join(BASE, "company.yaml")
KB = os.path.join(BASE, "kb")

from docx import Document

doc = Document()
doc.add_heading("某医院信息化运维服务项目投标书", level=0)
doc.add_heading("一、公司简介", level=1)
doc.add_paragraph("我公司成立于2018年，专注于医疗行业信息化运维，拥有ISO9001与软件企业认定资质。")
doc.add_heading("二、类似业绩", level=1)
doc.add_paragraph("近三年完成三甲医院运维项目5个，客户满意度98%。")
doc.add_heading("三、技术方案", level=1)
doc.add_paragraph("采用7x24小时驻场+远程双保障模式，建立巡检-预警-处置闭环。")
doc.add_heading("四、报价一览表", level=1)
t = doc.add_table(rows=2, cols=3)
t.rows[0].cells[0].text = "项目"
t.rows[0].cells[1].text = "数量"
t.rows[0].cells[2].text = "单价(万)"
t.rows[1].cells[0].text = "运维服务"
t.rows[1].cells[1].text = "1年"
t.rows[1].cells[2].text = "48"
doc.save(SAMPLE)

with open(TENDER, "w", encoding="utf-8") as f:
    f.write("招标公告：某医院信息化运维服务项目\n采购人：XX市第一人民医院\n预算：50万\n资质要求：营业执照、ISO9001\n评分：价格30分 技术40分 商务30分\n截止：2026-08-01\n")

os.makedirs(KB, exist_ok=True)
with open(COMPANY, "w", encoding="utf-8") as f:
    f.write("company: 示例科技有限公司\nfound_year: 2018\ncerts: [营业执照, ISO9001]\nsimilar_projects: 5\nbudget_hint: 50万\n")

print("sample files ready:", SAMPLE, TENDER, COMPANY)
