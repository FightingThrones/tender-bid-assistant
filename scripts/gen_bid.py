#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
招投标助手 - 标书生成引擎（catalog 版）。
输入: 招标公告文本 + 客户公司画像(yaml,可选) + 标书库 catalog
流程: 在 2300+ 真实模板中按行业+章节关键词匹配最相关模板 → 重新抽取该模板全文 →
      产出【推荐模板清单 + 标书大纲初稿 + 资质缺口 + 报价建议 + 风险提示】
用法:
  python gen_bid.py --tender 公告.txt [--company 客户画像.yaml] [--root "./private_bid_docs"]
"""
import os, sys, re, json, argparse, yaml
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import build_catalog as BC

ROOT_DEFAULT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "private_bid_docs"))
DATA = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "data")
DEFAULT_QUALS = ["营业执照", "法定代表人授权书", "类似业绩证明", "社保/纳税缴纳证明"]
RISK_PATTERNS = [
    ("陪标/串标红线", ["陪标", "串标", "围标", "控标", "内定", "返点", "回扣"], "BLOCK", "涉及违法违规投标表述，必须删除并停止该方向。"),
    ("真实性风险", ["伪造", "虚假", "借用资质", "挂靠", "假合同", "假业绩"], "BLOCK", "资质、业绩、授权必须真实可核验。"),
    ("高门槛资质", ["涉密", "保密资格", "等保测评", "安防资质", "电子与智能化", "厂商授权"], "REVIEW", "可能需要专项资质或原厂授权，投标前逐条核验。"),
    ("履约压力", ["7×24", "7*24", "驻场", "当天响应", "2小时", "应急"], "REVIEW", "需确认人员、排班、SLA 与成本是否支撑承诺。"),
]


# 通用投标流程词（几乎出现在每份标书里，对行业匹配无区分度，剔除）
DOMAIN_STOP = {
    "采购", "预算", "资质", "评分", "价格", "技术", "商务", "投标", "法人", "授权",
    "社保", "代理", "截止", "递交", "小微", "企业", "服务", "项目", "提供", "要求",
    "根据", "按照", "以下", "相关", "进行", "可以", "我们", "或者", "以及", "具有",
    "独立", "法定", "代表", "业绩", "考察", "重点", "建立", "安全", "备份", "响应",
    "小时", "数字化", "水平", "本次", "采用", "综合", "办法", "保障", "能力", "日常",
    "故障", "处置", "一套", "机制", "享受", "扣除", "优惠", "公告", "机构", "概况",
    "提升", "为提", "前递", "系统", "公司", "如下", "实施", "方案", "需求", "理解",
    "运维保", "一年", "人民币", "元万", "万元", "供货", "货物", "合同", "中标", "招标",
    "答疑", "澄清", "装订", "密封", "保证金", "函件", "承诺", "说明", "计划", "措施",
}


def extract_tokens(text):
    """抽取招标公告里的领域关键词（2-3字），剔除通用投标流程词与句子碎片。"""
    toks = set(re.findall(r"[\u4e00-\u9fff]{2,3}", text))
    return [t for t in toks if t not in DOMAIN_STOP and len(t) >= 2]


# 行业语义映射：哪些行业词代表「同一领域」，命中给大加权
DOMAIN_INDUSTRY = ["软件", "IT", "信息", "系统", "网络", "集成", "运维", "智能", "数据", "科技", "云计算", "互联网"]


def industry_boost(industry, scenario):
    blob = (industry or "") + " " + (scenario or "")
    if any(k in blob for k in DOMAIN_INDUSTRY):
        return 10
    return 0


def score_entry(entry, tokens, tender_text):
    blob = " ".join([entry["title"]] + entry["headings"] + [entry["snippet"]] +
                    [entry["industry"]] + [entry.get("scenario", "")])
    s = sum(1 for t in tokens if t in blob)
    s += industry_boost(entry["industry"], entry.get("scenario", ""))
    # 标题命中加权重
    for t in tokens:
        if t in entry["title"]:
            s += 2
    return s


def match(tender_text, catalog, top=3):
    tokens = extract_tokens(tender_text)
    scored = [(score_entry(e, tokens, tender_text), e) for e in catalog["entries"]]
    scored = [x for x in scored if x[0] > 0]
    scored.sort(key=lambda x: -x[0])
    return scored[:top], tokens


def extract_full(rel, root):
    full = os.path.join(root, rel)
    ext = rel.lower().split(".")[-1]
    if ext == "doc":
        _, hs, text = BC.extract_doc(full)
    elif ext == "docx":
        _, hs, text = BC.extract_docx(full)
    else:
        _, hs, text = BC.extract_pdf(full)
    return hs, text


def qual_gap(company):
    have = set(company.get("certs", [])) if company else set()
    return [q for q in DEFAULT_QUALS if q not in have]


def price_advice(tender_text, company):
    m = re.search(r"预算[：: ]*([\d.]+)\s*万", tender_text)
    budget = m.group(1) + "万" if m else (company or {}).get("budget_hint", "（见公告）")
    return f"报价建议参考预算 {budget}；注意是否含价格扣除/小微企业价格优惠，比对公告评分办法的价格分权重。"


def _first_match(patterns, text, default="（未识别，需人工补全）"):
    for pat in patterns:
        m = re.search(pat, text, re.S)
        if m:
            return re.sub(r"\s+", " ", m.group(1)).strip(" ：:\n\r\t。")
    return default


def tender_summary(tender_text):
    return {
        "project": _first_match([r"招标公告[：:]\s*([^\n]+)", r"项目名称[：:]\s*([^\n]+)", r"采购项目名称[：:]\s*([^\n]+)"], tender_text),
        "buyer": _first_match([r"采购人[：:]\s*([^\n]+)", r"招标人[：:]\s*([^\n]+)"], tender_text),
        "agency": _first_match([r"采购代理机构[：:]\s*([^\n]+)", r"代理机构[：:]\s*([^\n]+)"], tender_text),
        "budget": _first_match([r"(?:项目预算|预算金额|采购预算)[：:：\s人民币]*([\d.]+\s*万?元?)"], tender_text),
        "deadline": _first_match([r"(?:投标截止时间|截止时间|递交截止时间)[：:：\s]*([^\n。]+)"], tender_text),
        "method": _first_match([r"(?:评标|评分)采用([^。\n]+)", r"(综合评分法|最低评标价法|竞争性磋商|竞争性谈判|询价)"], tender_text),
    }


def required_quals(tender_text):
    checks = [
        ("营业执照", ["营业执照", "独立法人"]),
        ("法定代表人授权书", ["法定代表人授权", "法人授权"]),
        ("社保/纳税缴纳证明", ["社保", "纳税"]),
        ("类似业绩证明", ["类似业绩", "项目业绩", "近三年"]),
        ("软件企业/开发能力证明", ["软件企业", "软件开发资质", "系统开发"]),
        ("信息安全/等保相关能力", ["等保", "信息安全", "网络安全"]),
        ("厂商授权", ["厂商授权", "原厂授权"]),
    ]
    found = []
    for name, keys in checks:
        if any(k in tender_text for k in keys):
            found.append(name)
    return found or DEFAULT_QUALS


def qual_gap_from_tender(company, tender_text):
    have = set(company.get("certs", [])) if company else set()
    return [q for q in required_quals(tender_text) if q not in have]


def compliance_review(tender_text, company):
    blob = tender_text + "\n" + yaml.safe_dump(company or {}, allow_unicode=True)
    hits = []
    for title, keys, level, advice in RISK_PATTERNS:
        matched = [k for k in keys if k in blob]
        if matched:
            hits.append({"title": title, "level": level, "matched": matched, "advice": advice})
    if not hits:
        hits.append({"title": "合规初筛", "level": "PASS", "matched": [], "advice": "未发现明显红线词；仍需人工核验公告原文、资质真实性和报价依据。"})
    return hits


def response_matrix(tender_text):
    items = [
        ("资格条件", required_quals(tender_text), "商务标/资格审查文件"),
        ("价格评分", ["报价明细", "小微企业声明", "价格扣除说明"], "报价文件"),
        ("技术评分", ["需求理解", "总体方案", "实施计划", "运维保障", "数据安全", "应急响应"], "技术标"),
        ("服务承诺", ["服务期", "响应时间", "售后服务", "项目团队"], "技术标/商务标"),
    ]
    rows = []
    for category, points, section in items:
        for point in points:
            status = "需响应" if point in tender_text or category != "资格条件" else "建议核对"
            rows.append((category, point, status, section))
    return rows


def write_docx(markdown_text, out_path):
    try:
        import docx
    except Exception:
        return False, "未安装 python-docx，已跳过 Word 输出。"
    doc = docx.Document()
    for raw in markdown_text.splitlines():
        line = raw.strip()
        if not line:
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("> "):
            doc.add_paragraph(line[2:].strip())
        else:
            doc.add_paragraph(line)
    doc.save(out_path)
    return True, out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--tender", required=True)
    ap.add_argument("--company")
    ap.add_argument("--root", default=ROOT_DEFAULT)
    ap.add_argument("--catalog", default=os.path.join(DATA, "catalog.json"))
    ap.add_argument("--out", default=os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "output"))
    ap.add_argument("--top", type=int, default=3)
    ap.add_argument("--docx", action="store_true", help="同时输出 Word 版初稿")
    args = ap.parse_args()

    with open(args.tender, encoding="utf-8") as f:
        tender = f.read()
    company = None
    if args.company:
        with open(args.company, encoding="utf-8") as f:
            company = yaml.safe_load(f)

    with open(args.catalog, encoding="utf-8") as f:
        catalog = json.load(f)
    if not catalog.get("entries"):
        sys.exit("catalog 为空，请先运行 build_catalog.py。")

    ranked, tokens = match(tender, catalog, args.top)
    if not ranked:
        sys.exit("未匹配到任何模板，请检查公告文本或 catalog。")
    print(f"[匹配] 命中关键词 {len(tokens)} 个，返回 Top{len(ranked)} 模板：")
    for sc, e in ranked:
        print(f"  - {e['industry']} | {e['title']} | 分 {sc}")

    # 用最佳匹配模板抽取全文做大纲
    top_sc, top_e = ranked[0]
    hs, text = extract_full(top_e["rel"], args.root)
    outline = hs[:25] if hs else ["（该模板未识别到明确章节，已附全文快照）"]

    summary = tender_summary(tender)
    gaps = qual_gap_from_tender(company, tender)
    reviews = compliance_review(tender, company)
    matrix = response_matrix(tender)

    lines = []
    lines.append(f"# 投标文件初稿（基于最佳匹配模板：{top_e['industry']} / {top_e['title']}）\n")
    lines.append("> 本初稿由投标助手 skill 自动匹配真实模板生成，正式投递前须人工审核精修。\n")
    lines.append("## 一、招标要点摘要")
    lines.append(f"- 项目名称：{summary['project']}")
    lines.append(f"- 采购人：{summary['buyer']}")
    lines.append(f"- 代理机构：{summary['agency']}")
    lines.append(f"- 预算：{summary['budget']}")
    lines.append(f"- 截止时间：{summary['deadline']}")
    lines.append(f"- 评审方式：{summary['method']}\n")
    lines.append("## 二、合规红线与投标风险初筛")
    for item in reviews:
        matched = f"；命中：{', '.join(item['matched'])}" if item["matched"] else ""
        lines.append(f"- [{item['level']}] {item['title']}：{item['advice']}{matched}")
    lines.append("\n## 三、推荐使用的模板（按相关度）")
    for sc, e in ranked:
        lines.append(f"- 【{e['industry']}】{e['title']}  （相关度 {sc}）\n  路径：`{e['rel']}`")
    lines.append("\n## 四、标书大纲初稿（取自最佳匹配模板的章节结构）")
    for h in outline:
        lines.append(f"- {h}")
    lines.append("\n## 五、逐条响应矩阵（初版）")
    lines.append("| 类别 | 响应点 | 状态 | 建议放置位置 |")
    lines.append("| --- | --- | --- | --- |")
    for category, point, status, section in matrix:
        lines.append(f"| {category} | {point} | {status} | {section} |")
    lines.append("\n## 六、资质缺口清单（按公告关键词初筛）")
    for q in gaps:
        lines.append(f"- [建议准备] {q}")
    if not gaps:
        lines.append("- 暂未发现明显缺口；仍需按公告原文逐项核验。")
    lines.append("\n## 七、报价建议")
    lines.append(price_advice(tender, company))
    lines.append("\n## 八、商业交付清单")
    lines.append("- 公告要点摘要与投标可行性判断")
    lines.append("- 资质/业绩/人员/授权材料缺口清单")
    lines.append("- 商务标、技术标、报价文件初稿")
    lines.append("- 逐条响应矩阵与扣分风险提示")
    lines.append("- 最终投递前人工审核记录")
    lines.append("\n## 九、风险提示")
    lines.append("- 不同采购人评分标准差异大，务必逐条响应评分点；")
    lines.append("- 业绩/资质证明须真实可查，杜绝造假；")
    lines.append("- 初稿模板可能含与本项目无关的占位内容，须通篇替换。\n")

    os.makedirs(args.out, exist_ok=True)
    out_path = os.path.join(args.out, "标书初稿.md")
    md = "\n".join(lines)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(md)
    print("wrote", out_path)
    if args.docx:
        ok, info = write_docx(md, os.path.join(args.out, "标书初稿.docx"))
        print("wrote" if ok else "skip", info)


if __name__ == "__main__":
    main()
